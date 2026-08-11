import argparse
from docx import Document

def convert(md_file, output):
    doc = Document()
    with open(md_file, encoding="utf-8") as fh:
        for line in fh:
            line = line.rstrip()
            if line.startswith("# "):
                doc.add_heading(line[2:], level=0)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=1)
            elif line.startswith("### "):
                doc.add_heading(line[4:], level=2)
            elif line.startswith("- "):
                doc.add_paragraph(line[2:], style="List Bullet")
            elif line.startswith("> "):
                doc.add_paragraph(line[2:], style="Intense Quote")
            elif line.strip():
                doc.add_paragraph(line)
    doc.save(output)
    print(f"Wrote {output}")

def main():
    ap = argparse.ArgumentParser(description="Convert a markdown file to DOCX")
    ap.add_argument("input", help="input .md file")
    ap.add_argument("output", nargs="?", help="output .docx file (default: input.docx)")
    args = ap.parse_args()
    convert(args.input, args.output or args.input.rsplit(".", 1)[0] + ".docx")

if __name__ == "__main__":
    main()
